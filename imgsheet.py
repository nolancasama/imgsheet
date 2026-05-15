import os
import math
import json
import base64
import smtplib
import subprocess
import requests
from datetime import datetime
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from io import BytesIO
from PIL import Image, ImageFilter
from docx import Document
from docx.shared import Mm
import tkinter as tk
from tkinter import messagebox

try:
    import anthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False

# =========================
# CONFIG
# =========================
SERPAPI_KEY = os.environ.get("SERPAPI_KEY", "")
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")

GMAIL_USER     = os.environ.get("GMAIL_USER", "")
GMAIL_APP_PASS = os.environ.get("GMAIL_APP_PASS", "")

B4_WIDTH_MM = 257
B4_HEIGHT_MM = 364
ROWS = 5
COLS = 5
PAGE_MARGIN_MM = 5
PADDING_MM = 4

SHARPNESS_THRESHOLD      = 80    # higher = stricter blur rejection
BORDER_TOLERANCE         = 15    # lower = stricter solid-border rejection
VIBRANCE_THRESHOLD       = 0.12  # 0-1, min mean saturation (poster art needs color)
COMPRESSION_THRESHOLD    = 0.04  # bytes per pixel, below this = over-compressed
CENTER_BIAS_RATIO        = 0.35  # center edge density must be at least this fraction of corners

BLOCKED_DOMAINS = [
    "pinterest.com", "pinterest.co",
    "shutterstock.com", "gettyimages.com",
    "istockphoto.com", "alamy.com",
    "dreamstime.com", "123rf.com",
    "depositphotos.com", "stocksy.com",
]

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMP_DIR = os.path.join(BASE_DIR, "temp_images")
os.makedirs(TEMP_DIR, exist_ok=True)

HASHES_FILE    = os.path.join(BASE_DIR, "seen_hashes.json")
URLS_FILE      = os.path.join(BASE_DIR, "seen_urls.json")
RUN_COUNTS_FILE = os.path.join(BASE_DIR, "run_counts.json")

# =========================
# SESSION
# =========================
def create_session():
    session = requests.Session()
    retries = Retry(total=3, backoff_factor=1, status_forcelist=[403, 429, 500, 502, 503, 504])
    adapter = HTTPAdapter(max_retries=retries)
    session.mount("http://", adapter)
    session.mount("https://", adapter)
    session.headers.update({
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)",
        "Accept": "application/json,text/html;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
        "Referer": "https://www.google.com/"
    })
    return session

session = create_session()

# =========================
# SAFE WRAPPER
# =========================
def safe_run(func):
    def wrapper():
        try:
            func()
        except Exception as e:
            messagebox.showerror("Error", str(e))
    return wrapper

# =========================
# QUALITY FILTERS
# =========================
def sharpness_score(img):
    small = img.convert("L").resize((200, 200))
    edges = small.filter(ImageFilter.FIND_EDGES)
    pixels = list(edges.getdata())
    mean = sum(pixels) / len(pixels)
    return sum((p - mean) ** 2 for p in pixels) / len(pixels)

def has_solid_border(img, fraction=0.05):
    w, h = img.size
    bh = max(2, int(h * fraction))
    bw = max(2, int(w * fraction))
    bands = [
        img.crop((0, 0, w, bh)),
        img.crop((0, h - bh, w, h)),
        img.crop((0, 0, bw, h)),
        img.crop((w - bw, 0, w, h)),
    ]
    for band in bands:
        small = band.resize((20, 20))
        pixels = list(small.getdata())
        r = [p[0] for p in pixels]
        g = [p[1] for p in pixels]
        b = [p[2] for p in pixels]
        if (max(r) - min(r) < BORDER_TOLERANCE and
                max(g) - min(g) < BORDER_TOLERANCE and
                max(b) - min(b) < BORDER_TOLERANCE):
            return True
    return False

def color_vibrance(img):
    small = img.resize((100, 100))
    pixels = list(small.getdata())
    sats = []
    for r, g, b in pixels:
        mx = max(r, g, b)
        if mx > 0:
            sats.append((mx - min(r, g, b)) / mx)
    return sum(sats) / len(sats) if sats else 0

def is_subject_centered(img):
    w, h = img.size
    cx, cy = w // 3, h // 3
    def edge_density(region):
        pixels = list(region.convert("L").filter(ImageFilter.FIND_EDGES).getdata())
        return sum(pixels) / len(pixels) if pixels else 0
    center_density = edge_density(img.crop((cx, cy, cx * 2, cy * 2)))
    corner_density = sum(edge_density(img.crop(box)) for box in [
        (0, 0, cx, cy), (cx*2, 0, w, cy),
        (0, cy*2, cx, h), (cx*2, cy*2, w, h)
    ]) / 4
    return center_density >= corner_density * CENTER_BIAS_RATIO

def compression_ok(content_len, img):
    return (content_len / (img.width * img.height)) >= COMPRESSION_THRESHOLD

def is_blocked_domain(url):
    return any(domain in url for domain in BLOCKED_DOMAINS)

def claude_vision_check(img, prompt):
    try:
        key = ANTHROPIC_API_KEY or os.environ.get("ANTHROPIC_API_KEY", "")
        if not key:
            return True
        buffer = BytesIO()
        img.save(buffer, format="JPEG", quality=85)
        img_b64 = base64.b64encode(buffer.getvalue()).decode()
        client = anthropic.Anthropic(api_key=key)
        msg = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=20,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": img_b64}},
                    {"type": "text", "text": (
                        f"This image is a candidate for a poster, t-shirt, or trading card featuring: '{prompt}'.\n\n"
                        f"REJECT if ANY of the following are true:\n"
                        f"- The subject is not clearly visible or not the dominant focus\n"
                        f"- This is a photo of physical artwork, a print, poster, or screen showing the image\n"
                        f"- This is a photo of merchandise (figurine, box, product on a shelf)\n"
                        f"- The image contains multiple panels, comparisons, or collages\n"
                        f"- There are watermarks, signatures, or overlaid text (beyond minor corner signatures)\n"
                        f"- The subject occupies less than 30% of the image\n"
                        f"- The art quality looks amateur, heavily pixelated, or like a low-effort edit\n\n"
                        f"ACCEPT only if it is clean, single-subject artwork or illustration that would look great printed on merchandise.\n\n"
                        f"Reply with ACCEPT or REJECT followed by a 3-word reason."
                    )}
                ]
            }]
        )
        result = msg.content[0].text.strip().upper()
        print(f"Claude: {msg.content[0].text.strip()}")
        return result.startswith("ACCEPT")
    except Exception as e:
        print("Claude vision error:", e)
        return True

# =========================
# HASH
# =========================
def average_hash(img, size=8):
    img = img.convert("L").resize((size, size))
    pixels = list(img.getdata())
    avg = sum(pixels) / len(pixels)
    return tuple(p > avg for p in pixels)

# =========================
# SEARCH
# =========================
def search_images(query, count, start=0):
    urls = []
    try:
        res = requests.get(
            "https://serpapi.com/search",
            params={
                "engine": "google_images",
                "q": query,
                "api_key": SERPAPI_KEY,
                "num": count,
                "start": start,
                "imgsz": "l",   # large images only
                "imgar": "t",   # tall/portrait aspect ratio
            },
            timeout=10
        )
        data = res.json()
        for item in data.get("images_results", []):
            url = item.get("original")
            if (url
                    and item.get("original_width", 0) >= 400
                    and item.get("original_height", 0) >= 400
                    and not is_blocked_domain(url)):
                urls.append(url)
    except Exception as e:
        print("Search error:", e)
    print(f"Found {len(urls)} URLs")
    return urls

# =========================
# IMAGE PROCESS
# =========================
def process_image(url, target_px, seen):
    try:
        res = session.get(url, timeout=10)
        if res.status_code != 200:
            return None
        if "image" not in res.headers.get("Content-Type", ""):
            return None

        img = Image.open(BytesIO(res.content)).convert("RGB")

        if img.width < 400 or img.height < 400:
            return None

        if not compression_ok(len(res.content), img):
            print("Rejected: over-compressed")
            return None

        if sharpness_score(img) < SHARPNESS_THRESHOLD:
            print("Rejected: blurry")
            return None

        if has_solid_border(img):
            print("Rejected: border/frame")
            return None

        if color_vibrance(img) < VIBRANCE_THRESHOLD:
            print("Rejected: low color vibrance")
            return None

        if not is_subject_centered(img):
            print("Rejected: subject not centered")
            return None

        h = average_hash(img)
        if h in seen:
            return None
        seen.add(h)

        img.thumbnail(target_px)
        bg = Image.new("RGB", target_px, (255, 255, 255))
        x = (target_px[0] - img.width) // 2
        y = (target_px[1] - img.height) // 2
        bg.paste(img, (x, y))
        return bg

    except:
        return None

# =========================
# DOC
# =========================
def create_doc(image_paths, output_name):
    doc = Document()
    section = doc.sections[0]
    section.page_width  = Mm(B4_WIDTH_MM)
    section.page_height = Mm(B4_HEIGHT_MM)
    section.top_margin    = Mm(PAGE_MARGIN_MM)
    section.bottom_margin = Mm(PAGE_MARGIN_MM)
    section.left_margin   = Mm(PAGE_MARGIN_MM)
    section.right_margin  = Mm(PAGE_MARGIN_MM)

    table = doc.add_table(rows=ROWS, cols=COLS)
    usable_w = B4_WIDTH_MM - (PAGE_MARGIN_MM * 2)
    img_w = (usable_w / COLS) - PADDING_MM

    i = 0
    for r in range(ROWS):
        for c in range(COLS):
            if i >= len(image_paths):
                break
            cell = table.cell(r, c)
            para = cell.paragraphs[0]
            para.alignment = 1
            para.add_run().add_picture(image_paths[i], width=Mm(img_w))
            i += 1

    doc.save(output_name)

# =========================
# EMAIL
# =========================
def send_email(to_addr, file_path):
    msg = MIMEMultipart()
    msg["From"]    = GMAIL_USER
    msg["To"]      = to_addr
    msg["Subject"] = f"Image Sheet — {os.path.basename(file_path)}"
    msg.attach(MIMEText("Your image sheet is attached.", "plain"))

    with open(file_path, "rb") as f:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(f.read())
    encoders.encode_base64(part)
    part.add_header("Content-Disposition", f'attachment; filename="{os.path.basename(file_path)}"')
    msg.attach(part)

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(GMAIL_USER, GMAIL_APP_PASS)
        server.send_message(msg)

# =========================
# MAIN
# =========================
@safe_run
def generate():
    status_label.config(text="Working...")

    raw = prompt_box.get("1.0", tk.END)
    prompts = [p.strip() for p in raw.replace(",", "\n").splitlines() if p.strip()]
    if not prompts:
        messagebox.showerror("Error", "Enter at least one prompt")
        return
    if len(prompts) > ROWS * COLS:
        messagebox.showerror("Error", f"Max {ROWS * COLS} prompts.")
        return

    use_claude = claude_vision_var.get()
    if use_claude and not ANTHROPIC_AVAILABLE:
        messagebox.showerror("Error", "anthropic package not installed.\nRun: pip install anthropic")
        return
    if use_claude and not (ANTHROPIC_API_KEY or os.environ.get("ANTHROPIC_API_KEY")):
        messagebox.showerror("Error", "Set ANTHROPIC_API_KEY in the script or as an environment variable.")
        return

    track = avoid_dupes_var.get()
    seen = set()
    used_urls = set()
    run_counts = {}

    if track:
        if os.path.exists(HASHES_FILE):
            with open(HASHES_FILE) as f:
                seen = set(tuple(h) for h in json.load(f))
        if os.path.exists(URLS_FILE):
            with open(URLS_FILE) as f:
                used_urls = set(json.load(f))

    if os.path.exists(RUN_COUNTS_FILE):
        with open(RUN_COUNTS_FILE) as f:
            run_counts = json.load(f)

    total = ROWS * COLS
    base, extra = divmod(total, len(prompts))
    prompt_counts = [base + 1 if i < extra else base for i in range(len(prompts))]
    all_images = []

    for i, prompt in enumerate(prompts):
        print("\n=== PROMPT:", prompt, "===")
        run_num = run_counts.get(prompt, 0)
        needed  = prompt_counts[i]
        urls    = search_images(prompt, needed * 6, start=run_num * 100)

        if track:
            urls = [u for u in urls if u not in used_urls]

        count = 0
        for url in urls:
            if count >= needed:
                break

            img = process_image(url, (700, 1000), seen)
            if not img:
                continue

            if use_claude:
                status_label.config(text=f"Claude checking image {len(all_images)+1}...")
                root.update_idletasks()
                if not claude_vision_check(img, prompt):
                    print("Rejected by Claude Vision")
                    continue

            path = os.path.join(TEMP_DIR, f"{len(all_images)}.jpg")
            img.save(path)
            all_images.append(path)
            used_urls.add(url)
            count += 1

        run_counts[prompt] = run_num + 1

    print("TOTAL IMAGES:", len(all_images))

    if track:
        with open(HASHES_FILE, "w") as f:
            json.dump([list(h) for h in seen], f)
        with open(URLS_FILE, "w") as f:
            json.dump(list(used_urls), f)

    with open(RUN_COUNTS_FILE, "w") as f:
        json.dump(run_counts, f)

    timestamp   = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_name = os.path.join(BASE_DIR, f"output_{timestamp}.docx")
    create_doc(all_images, output_name)

    if pdf_var.get():
        soffice_candidates = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            r"C:\Program Files\OpenOffice 4\program\soffice.exe",
            r"C:\Program Files (x86)\OpenOffice 4\program\soffice.exe",
            r"C:\Program Files\OpenOffice.org 3\program\soffice.exe",
            r"C:\Program Files (x86)\OpenOffice.org 3\program\soffice.exe",
        ]
        soffice = next((p for p in soffice_candidates if os.path.exists(p)), None)
        if not soffice:
            messagebox.showwarning("PDF Failed", "OpenOffice/LibreOffice not found.")
            status_label.config(text=f"Done! Saved as {os.path.basename(output_name)}")
        else:
            status_label.config(text="Converting to PDF...")
            try:
                subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir", BASE_DIR, output_name], check=True)
                pdf_name = os.path.basename(output_name).replace(".docx", ".pdf")
                status_label.config(text=f"Done! {os.path.basename(output_name)} + {pdf_name}")
                messagebox.showinfo("Done", f"Sheet created!\n{os.path.basename(output_name)}\n{pdf_name}")
            except Exception as e:
                status_label.config(text="Done! PDF conversion failed.")
                messagebox.showwarning("PDF Failed", f"Saved docx but PDF failed:\n{e}")
    else:
        status_label.config(text=f"Done! Saved as {os.path.basename(output_name)}")
        messagebox.showinfo("Done", f"Sheet created!\n{os.path.basename(output_name)}")

    if email_var.get():
        to_addr = email_entry.get().strip()
        if not to_addr:
            messagebox.showwarning("Email", "Enter a recipient email address.")
        elif not GMAIL_USER or not GMAIL_APP_PASS:
            messagebox.showwarning("Email", "Set GMAIL_USER and GMAIL_APP_PASS in the script.")
        else:
            try:
                status_label.config(text="Sending email...")
                root.update_idletasks()
                send_email(to_addr, output_name)
                status_label.config(text="Done! Email sent.")
                messagebox.showinfo("Email Sent", f"Sheet sent to {to_addr}")
            except Exception as e:
                messagebox.showerror("Email Failed", str(e))

# =========================
# UI
# =========================
root = tk.Tk()
root.title("Image Sheet Generator")
root.geometry("500x460")

tk.Label(root, text="Enter up to 25 prompts (comma or newline separated):").pack(pady=5)

prompt_box = tk.Text(root, width=50, height=8)
prompt_box.pack(pady=2)

pdf_var = tk.BooleanVar()
tk.Checkbutton(root, text="Also export as PDF (requires OpenOffice/LibreOffice)", variable=pdf_var).pack()

avoid_dupes_var = tk.BooleanVar()
tk.Checkbutton(root, text="Avoid previously used images", variable=avoid_dupes_var).pack()

claude_vision_var = tk.BooleanVar()
tk.Checkbutton(root, text="Filter with Claude Vision (requires ANTHROPIC_API_KEY)", variable=claude_vision_var).pack()

email_var = tk.BooleanVar()
tk.Checkbutton(root, text="Send to email:", variable=email_var).pack()
email_entry = tk.Entry(root, width=40)
email_entry.pack(pady=2)

tk.Button(root, text="Generate", command=generate).pack(pady=10)

status_label = tk.Label(root, text="Ready")
status_label.pack()

root.mainloop()
