import os
import re
import math
import random
import smtplib
import subprocess
import threading
import requests
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from typing import Callable, Optional
from datetime import datetime
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from io import BytesIO
from PIL import Image, ImageFilter, ImageFile
from docx import Document
from docx.shared import Mm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Allow PIL to load slightly-truncated downloads instead of erroring them out
ImageFile.LOAD_TRUNCATED_IMAGES = True

# =========================
# CONFIG
# =========================
SERPAPI_KEY = os.environ.get("SERPAPI_KEY", "")
BRAVE_API_KEY = os.environ.get("BRAVE_API_KEY", "")
SERPER_API_KEY = os.environ.get("SERPER_API_KEY", "")
GMAIL_USER = os.environ.get("GMAIL_USER", "")
GMAIL_APP_PASS = os.environ.get("GMAIL_APP_PASS", "")

B4_WIDTH_MM = 257
B4_HEIGHT_MM = 364
ROWS = 5
COLS = 5
PAGE_MARGIN_MM = 5
PADDING_MM = 4

SHARPNESS_THRESHOLD = 80
BORDER_TOLERANCE = 15
VIBRANCE_THRESHOLD = 0.07
COMPRESSION_THRESHOLD = 0.04
CENTER_BIAS_RATIO = 0.35
MAX_CANDIDATES_PER_PROMPT = 15

PAPER_SIZES = {
    "B4":     (257, 364),
    "B5":     (182, 257),
    "A4":     (210, 297),
    "Letter": (216, 279),
    "A3":     (297, 420),
}

BLOCKED_DOMAINS = [
    "pinterest.com", "pinterest.co",
    "shutterstock.com", "gettyimages.com",
    "istockphoto.com", "alamy.com",
    "dreamstime.com", "123rf.com",
    "depositphotos.com", "stocksy.com",
    "amazon.com", "amazon.co.jp", "ebay.com", "etsy.com",
    "amiami.com", "goodsmile.info", "solarisjapan.com",
    "hlj.com", "bigbadtoystore.com", "animegami.com",
    "play-asia.com", "walmart.com", "target.com",
]

MERCHANDISE_KEYWORDS = [
    "buy", "shop", "sale", "sold", "price", "usd", "\\$",
    "figurine", "figure", "plush", "toy", "doll", "statue",
    "merchandise", "merch", "product", "item", "listing",
    "stock", "order", "shipping", "cart", "checkout",
    "amazon", "ebay", "etsy", "walmart", "target",
    "screenshot", "screencap", "screen shot",
    "chapter", "volume", "scanlation", "raw scan",
    "tweet", "tweeted", "reddit", "discord message",
]

BLOCKED_LINK_DOMAINS = [
    "reddit.com", "i.redd.it",
    "twitter.com/status",
    "mangadex.org", "mangakakalot", "manganelo",
]

HERO_DOMAINS = [
    "static.wikia.nocookie.net", "fandom.com",
    "cdn.myanimelist.net", "myanimelist.net",
    "zerochan.net", "safebooru.org", "danbooru.donmai.us",
    "artstation.com", "cdn.pixiv.net", "i.pximg.net",
]

HERO_KEYWORDS = [
    "wiki", "character", "artwork", "render", "official",
    "transparent", "png", "key visual", "fanart", "illustration",
]

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMP_DIR = os.path.join(BASE_DIR, "temp_images")
os.makedirs(TEMP_DIR, exist_ok=True)

STYLE_MODIFIERS = [
    "official art", "fanart", "illustration",
    "key visual", "anime art", "digital art",
    "concept art", "character art",
]


# =========================
# DATA CLASSES
# =========================
@dataclass
class PipelineOptions:
    export_pdf: bool = False
    randomize: bool = False
    search_engine: str = "serpapi"
    send_email_to: str = ""
    rows: int = 5
    cols: int = 5
    paper_size: str = "B4"
    double_sided: bool = False


@dataclass
class CharacterResult:
    prompt: str
    image_paths: list        # selected images for the sheet
    candidate_paths: list = field(default_factory=list)  # all passing candidates


@dataclass
class PipelineResult:
    output_docx: str
    output_pdf: Optional[str]
    characters: list  # list of CharacterResult
    back_image_paths: list = field(default_factory=list)


# =========================
# SESSION (module-level)
# =========================
def create_session():
    s = requests.Session()
    # This session is used only for image downloads. Keep retries minimal: a 403
    # (hotlink-blocked) never recovers on retry, and long backoffs stall the
    # 4-worker download pool on dead URLs — the main cause of slow generation.
    retries = Retry(total=1, backoff_factor=0.3, status_forcelist=[429, 503])
    adapter = HTTPAdapter(max_retries=retries)
    s.mount("http://", adapter)
    s.mount("https://", adapter)
    s.headers.update({
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)",
        "Accept": "application/json,text/html;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
        "Referer": "https://www.google.com/"
    })
    return s


session = create_session()


# =========================
# CANCELLATION / PARTIAL RESULTS
# =========================
class PipelineCancelled(Exception):
    def __init__(self, partial_images, character_results):
        super().__init__("Cancelled.")
        self.partial_images = partial_images
        self.character_results = character_results


def clear_temp_dir():
    for fname in os.listdir(TEMP_DIR):
        try:
            os.remove(os.path.join(TEMP_DIR, fname))
        except Exception:
            pass


def fill_to_count_spread(character_results, total):
    """Fill image list to total by round-robining across characters."""
    groups = [list(cr.image_paths) for cr in character_results if cr.image_paths]
    if not groups:
        return []
    result = []
    i = 0
    while len(result) < total:
        char_idx = i % len(groups)
        img_idx = i // len(groups)
        result.append(groups[char_idx][img_idx % len(groups[char_idx])])
        i += 1
    return result[:total]


# =========================
# QUALITY FILTERS
# =========================
def sharpness_score(img):
    small = img.convert("L").resize((200, 200))
    edges = small.filter(ImageFilter.FIND_EDGES)
    pixels = list(edges.getdata())
    mean = sum(pixels) / len(pixels)
    return sum((p - mean) ** 2 for p in pixels) / len(pixels)


def has_solid_border(img, fraction=0.05):
    w, h = img.size
    bh = max(2, int(h * fraction))
    bw = max(2, int(w * fraction))
    bands = [
        img.crop((0, 0, w, bh)),
        img.crop((0, h - bh, w, h)),
        img.crop((0, 0, bw, h)),
        img.crop((w - bw, 0, w, h)),
    ]
    for band in bands:
        small = band.resize((20, 20))
        pixels = list(small.getdata())
        r = [p[0] for p in pixels]
        g = [p[1] for p in pixels]
        b = [p[2] for p in pixels]
        if (max(r) - min(r) < BORDER_TOLERANCE and
                max(g) - min(g) < BORDER_TOLERANCE and
                max(b) - min(b) < BORDER_TOLERANCE):
            return True
    return False


def color_vibrance(img):
    small = img.resize((100, 100))
    pixels = list(small.getdata())
    sats = []
    for r, g, b in pixels:
        mx = max(r, g, b)
        if mx > 0:
            sats.append((mx - min(r, g, b)) / mx)
    return sum(sats) / len(sats) if sats else 0


def is_subject_centered(img):
    w, h = img.size
    cx, cy = w // 3, h // 3

    def edge_density(region):
        pixels = list(region.convert("L").filter(ImageFilter.FIND_EDGES).getdata())
        return sum(pixels) / len(pixels) if pixels else 0

    center_density = edge_density(img.crop((cx, cy, cx * 2, cy * 2)))
    corner_density = sum(edge_density(img.crop(box)) for box in [
        (0, 0, cx, cy), (cx * 2, 0, w, cy),
        (0, cy * 2, cx, h), (cx * 2, cy * 2, w, h)
    ]) / 4
    return center_density >= corner_density * CENTER_BIAS_RATIO


def compression_ok(content_len, img):
    return (content_len / (img.width * img.height)) >= COMPRESSION_THRESHOLD


def is_blocked_domain(url):
    return any(domain in url for domain in BLOCKED_DOMAINS)


def hero_score(item):
    url = item.get("original", "")
    meta = " ".join(filter(None, [
        item.get("title", ""),
        item.get("source", ""),
        item.get("link", ""),
        url,
    ])).lower()
    score = 0
    if any(d in url for d in HERO_DOMAINS):
        score += 3
    score += sum(1 for kw in HERO_KEYWORDS if kw in meta)
    return score


def smart_crop(img):
    """Crop to the most visually active portrait-aspect region."""
    w, h = img.size
    target_ratio = 700 / 1000  # portrait card ratio

    if w / h > target_ratio:
        crop_h = h
        crop_w = int(h * target_ratio)
    else:
        crop_w = w
        crop_h = min(h, int(w / target_ratio))

    # Skip if the crop is already nearly the full image
    if crop_w >= w * 0.92 and crop_h >= h * 0.92:
        return img

    steps = 8
    step_x = max(1, (w - crop_w) // steps)
    step_y = max(1, (h - crop_h) // steps)

    best_score = -1
    best_box = ((w - crop_w) // 2, (h - crop_h) // 2,
                (w + crop_w) // 2, (h + crop_h) // 2)

    gray = img.convert("L")
    for y in range(0, max(1, h - crop_h + 1), step_y):
        for x in range(0, max(1, w - crop_w + 1), step_x):
            box = (x, y, x + crop_w, y + crop_h)
            region = gray.crop(box).resize((40, 60)).filter(ImageFilter.FIND_EDGES)
            score = sum(region.getdata()) / (40 * 60)
            if score > best_score:
                best_score = score
                best_box = box

    return img.crop(best_box)


def average_hash(img, size=8):
    img = img.convert("L").resize((size, size))
    pixels = list(img.getdata())
    avg = sum(pixels) / len(pixels)
    return tuple(p > avg for p in pixels)


# =========================
# SEARCH
# =========================
def _filter_and_rank(raw_items):
    candidates = []
    for item in raw_items:
        url = item.get("original")
        if not url:
            continue
        w = item.get("original_width", 0)
        h = item.get("original_height", 0)
        if w > 0 and h > 0 and (w < 400 or h < 400):
            continue
        if is_blocked_domain(url):
            continue
        meta = " ".join(filter(None, [
            item.get("title", ""),
            item.get("source", ""),
            item.get("link", ""),
        ])).lower()
        if any(kw in meta for kw in MERCHANDISE_KEYWORDS):
            continue
        if any(d in item.get("link", "") for d in BLOCKED_LINK_DOMAINS):
            continue
        candidates.append((hero_score(item), url))
    candidates.sort(key=lambda x: x[0], reverse=True)
    return [url for _, url in candidates]


def _search_brave(query, count, start=0):
    try:
        res = requests.get(
            "https://api.search.brave.com/res/v1/images/search",
            headers={
                "Accept": "application/json",
                "Accept-Encoding": "gzip",
                "X-Subscription-Token": BRAVE_API_KEY,
            },
            params={
                "q": query,
                "count": min(count, 100),
                "offset": start,
                "safesearch": "off",
            },
            timeout=20
        )
        if res.status_code != 200:
            print(f"Brave search error: HTTP {res.status_code} — {res.text[:200]}")
            return []
        items = []
        for r in res.json().get("results", []):
            props = r.get("properties", {})
            items.append({
                "original": props.get("url") or r.get("url"),
                "original_width": 0,
                "original_height": 0,
                "title": r.get("title", ""),
                "source": r.get("source", ""),
                "link": r.get("url", ""),
            })
        return items
    except Exception as e:
        print("Brave search error:", e)
        return []


def _search_serpapi(query, count, start=0):
    try:
        res = requests.get(
            "https://serpapi.com/search",
            params={
                "engine": "google_images",
                "q": query,
                "api_key": SERPAPI_KEY,
                "num": count,
                "start": start,
                "imgsz": "l",
                "imgar": "t",
            },
            timeout=20
        )
        if res.status_code != 200:
            print(f"SerpAPI search error: HTTP {res.status_code} — {res.text[:200]}")
            return []
        data = res.json()
        if data.get("error"):
            print(f"SerpAPI error: {data['error']}")
            return []
        items = []
        for r in data.get("images_results", []):
            items.append({
                "original": r.get("original"),
                "original_width": r.get("original_width", 0),
                "original_height": r.get("original_height", 0),
                "title": r.get("title", ""),
                "source": r.get("source", ""),
                "link": r.get("link", ""),
            })
        return items
    except Exception as e:
        print("SerpAPI search error:", e)
        return []


def _search_serper(query, count, start=0):
    try:
        res = requests.post(
            "https://google.serper.dev/images",
            headers={"X-API-KEY": SERPER_API_KEY, "Content-Type": "application/json"},
            json={"q": query, "num": min(count, 100), "start": start},
            timeout=20
        )
        if res.status_code != 200:
            print(f"Serper search error: HTTP {res.status_code} — {res.text[:200]}")
            return []
        items = []
        for r in res.json().get("images", []):
            items.append({
                "original": r.get("imageUrl"),
                "original_width": r.get("imageWidth", 0),
                "original_height": r.get("imageHeight", 0),
                "title": r.get("title", ""),
                "source": r.get("source", ""),
                "link": r.get("link", ""),
            })
        return items
    except Exception as e:
        print("Serper search error:", e)
        return []


def _search_ddg(query, count, start=0):
    try:
        from ddgs import DDGS
        items = []
        ddgs = DDGS(
            headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"},
            timeout=20,
        )
        for r in ddgs.images(query, max_results=count + start):
            items.append({
                "original": r.get("image"),
                "original_width": r.get("width", 0),
                "original_height": r.get("height", 0),
                "title": r.get("title", ""),
                "source": r.get("source", ""),
                "link": r.get("url", ""),
            })
        return items[start:]
    except Exception as e:
        print("DuckDuckGo search error:", e)
        return []


def search_images(query, count, start=0, engine="serpapi"):
    if engine == "serpapi":
        raw = _search_serpapi(query, count, start)
    elif engine == "brave":
        raw = _search_brave(query, count, start)
    elif engine == "serper":
        raw = _search_serper(query, count, start)
    elif engine == "test":
        return []  # test mode skips URL search entirely
    else:
        raw = _search_ddg(query, count, start)
    urls = _filter_and_rank(raw)
    print(f"Found {len(urls)} URLs for query: {query}")
    return urls


# =========================
# IMAGE PROCESS
# =========================
def process_image(url, target_px, seen, seen_lock=None, save_dir=None):
    try:
        res = session.get(url, timeout=6)
        if res.status_code != 200:
            return None
        if "image" not in res.headers.get("Content-Type", ""):
            return None

        raw = Image.open(BytesIO(res.content))
        if raw.mode == "P" and "transparency" in raw.info:
            raw = raw.convert("RGBA")
        img = raw.convert("RGB")
        raw = None  # free raw immediately

        if img.width < 400 or img.height < 400:
            return None

        if sharpness_score(img) < SHARPNESS_THRESHOLD:
            print("Rejected: blurry")
            return None

        h = average_hash(img)
        if seen_lock:
            with seen_lock:
                if h in seen:
                    return None
                seen.add(h)
        else:
            if h in seen:
                return None
            seen.add(h)

        img.thumbnail((1400, 2000))
        img = smart_crop(img)
        img.thumbnail(target_px)
        bg = Image.new("RGB", target_px, (255, 255, 255))
        x = (target_px[0] - img.width) // 2
        y = (target_px[1] - img.height) // 2
        bg.paste(img, (x, y))
        img = None  # free before saving

        if save_dir:
            import uuid
            path = os.path.join(save_dir, f"img_{uuid.uuid4().hex[:10]}.jpg")
            bg.save(path, "JPEG", quality=90)
            bg = None
            return path
        return bg

    except Exception as e:
        print(f"process_image dropped {url[:80]} ({type(e).__name__})")
        return None


# =========================
# DOC
# =========================
# Reference cell size derived from B4 5x5 — fixed across all paper sizes
_B4_USABLE_W = PAPER_SIZES["B4"][0] - PAGE_MARGIN_MM * 2
FIXED_CELL_W_MM = (_B4_USABLE_W / 5) - PADDING_MM  # 45.4mm


def _mirror_rows(paths, rows, cols):
    """Return paths with each row's columns reversed (for back-side of double-sided sheet)."""
    result = []
    for r in range(rows):
        row = paths[r * cols:(r + 1) * cols]
        result.extend(reversed(row))
    return result


def _build_table(doc, image_paths, rows, cols):
    table = doc.add_table(rows=rows, cols=cols)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'CCCCCC')
        tblBorders.append(border)
    table._tbl.tblPr.append(tblBorders)
    i = 0
    for r in range(rows):
        for c in range(cols):
            if i >= len(image_paths):
                break
            cell = table.cell(r, c)
            para = cell.paragraphs[0]
            para.alignment = 1
            try:
                para.add_run().add_picture(image_paths[i], width=Mm(FIXED_CELL_W_MM))
            except Exception as e:
                # A single unreadable image shouldn't sink the whole document —
                # leave the cell blank and keep going.
                print(f"Skipping unreadable image {image_paths[i]}: {e}")
            i += 1


def create_doc(image_paths, output_name, rows=5, cols=5, paper_size="B4", back_paths=None):
    w_mm, h_mm = PAPER_SIZES.get(paper_size, PAPER_SIZES["B4"])
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(w_mm)
    section.page_height = Mm(h_mm)

    # Centre the table by computing side margins so cells stay at fixed size
    grid_w = cols * (FIXED_CELL_W_MM + PADDING_MM)
    side_margin = max(PAGE_MARGIN_MM, (w_mm - grid_w) / 2)
    section.top_margin = Mm(PAGE_MARGIN_MM)
    section.bottom_margin = Mm(PAGE_MARGIN_MM)
    section.left_margin = Mm(side_margin)
    section.right_margin = Mm(side_margin)

    _build_table(doc, image_paths, rows, cols)

    if back_paths:
        # Page break then back sheet (mirrored columns) in the same document
        para = doc.add_paragraph()
        run = para.add_run()
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        run._r.append(br)
        _build_table(doc, back_paths, rows, cols)

    doc.save(output_name)


# =========================
# EMAIL
# =========================
def send_email(to_addr, file_path):
    msg = MIMEMultipart()
    msg["From"] = GMAIL_USER
    msg["To"] = to_addr
    msg["Subject"] = f"Image Sheet — {os.path.basename(file_path)}"
    msg.attach(MIMEText("Your image sheet is attached.", "plain"))

    with open(file_path, "rb") as f:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(f.read())
    encoders.encode_base64(part)
    part.add_header("Content-Disposition", f'attachment; filename="{os.path.basename(file_path)}"')
    msg.attach(part)

    with smtplib.SMTP_SSL("smtp.gmail.com", 465, timeout=15) as server:
        server.login(GMAIL_USER, GMAIL_APP_PASS)
        server.send_message(msg)


# =========================
# FIND SOFFICE
# =========================
def find_soffice():
    candidates = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        r"C:\Program Files\OpenOffice 4\program\soffice.exe",
        r"C:\Program Files (x86)\OpenOffice 4\program\soffice.exe",
        r"C:\Program Files\OpenOffice.org 3\program\soffice.exe",
        r"C:\Program Files (x86)\OpenOffice.org 3\program\soffice.exe",
    ]
    return next((p for p in candidates if os.path.exists(p)), None)


# =========================
# MAIN PIPELINE
# =========================
def make_sheet_filename(prompts, timestamp, suffix=""):
    joined = "_".join(prompts)
    safe = re.sub(r'[^a-zA-Z0-9]', '_', joined)
    safe = re.sub(r'_+', '_', safe).strip('_')
    if len(safe) > 60:
        safe = safe[:60].rstrip('_')
    return f"{safe}_{timestamp}{suffix}.docx"


def run_pipeline(prompts: list, options: PipelineOptions, on_progress: Callable[[str], None], cancel_event=None, output_dir=None, temp_dir=None) -> PipelineResult:
    if temp_dir is None:
        temp_dir = TEMP_DIR
        clear_temp_dir()
    else:
        os.makedirs(temp_dir, exist_ok=True)

    if not prompts:
        raise ValueError("Enter at least one prompt.")
    if len(prompts) > ROWS * COLS:
        raise ValueError(f"Max {ROWS * COLS} prompts.")

    seen = set()

    total = options.rows * options.cols
    base, extra = divmod(total, len(prompts))
    prompt_counts = [base + 1 if i < extra else base for i in range(len(prompts))]

    all_images = []
    all_back_images = []
    character_results = []
    dry_prompts = set()  # prompts that yielded no usable candidates — skip in shortfall
    seen_lock = threading.Lock()

    # Parallel search phase — all prompts searched simultaneously
    def _search(args):
        i, prompt, needed = args
        start = random.randint(0, 14) if options.randomize else 0
        print(f"\n=== PROMPT: {prompt} | start: {start} ===")
        urls = search_images(prompt, needed * 6, start=start, engine=options.search_engine)
        return i, prompt, needed, urls

    on_progress("Searching for all characters...")
    search_results = [None] * len(prompts)
    with ThreadPoolExecutor(max_workers=min(len(prompts), 5)) as sex:
        sfuts = {sex.submit(_search, (i, p, prompt_counts[i])): i for i, p in enumerate(prompts)}
        for sf in as_completed(sfuts):
            i, prompt, needed, urls = sf.result()
            search_results[i] = (prompt, needed, urls)
            on_progress(f"Found {len(urls)} URLs for: {prompt}")

    try:
        for i in range(len(prompts)):
            if cancel_event and cancel_event.is_set():
                raise InterruptedError("Cancelled by user.")

            prompt, needed, urls = search_results[i]
            char_candidate_paths = []
            char_paths = []
            on_progress(f"Processing: {prompt} ({len(urls)} candidates)...")

            # Test mode: generate solid-color placeholder images, no network
            if options.search_engine == "test":
                pool_size = needed * 3
                for j in range(pool_size):
                    h = hash(f"{prompt}{j}") & 0xFFFFFF
                    color = ((h >> 16) & 0xFF, (h >> 8) & 0xFF, h & 0xFF)
                    img = Image.new("RGB", (700, 1000), color)
                    path = os.path.join(temp_dir, f"c{i}_{j}.jpg")
                    img.save(path, "JPEG", quality=90)
                    char_candidate_paths.append(path)
                char_paths = char_candidate_paths[:needed]
                all_images.extend(char_paths)
                if options.double_sided:
                    char_back = char_candidate_paths[needed:needed * 2]
                    if len(char_back) < needed:
                        char_back = [char_paths[j % len(char_paths)] for j in range(needed)]
                    all_back_images.extend(char_back)
                character_results.append(CharacterResult(
                    prompt=prompt, image_paths=char_paths, candidate_paths=char_candidate_paths
                ))
                on_progress(f"{prompt}: {len(char_paths)} placeholder images")
                continue

            # Phase 1: parallel downloads — save every passing image to disk immediately
            with ThreadPoolExecutor(max_workers=4) as executor:
                futures = {executor.submit(process_image, url, (700, 1000), seen, seen_lock, temp_dir): url for url in urls}
                for future in as_completed(futures):
                    if cancel_event and cancel_event.is_set():
                        for f in futures:
                            f.cancel()
                        character_results.append(CharacterResult(
                            prompt=prompt, image_paths=[], candidate_paths=char_candidate_paths
                        ))
                        raise InterruptedError("Cancelled by user.")
                    path = future.result()
                    if path:
                        char_candidate_paths.append(path)
                    _cand_limit = max(MAX_CANDIDATES_PER_PROMPT, needed * 2) if options.double_sided else MAX_CANDIDATES_PER_PROMPT
                    if len(char_candidate_paths) >= _cand_limit:
                        for f in futures:
                            f.cancel()
                        break

            # Phase 2: take the first `needed` candidates for the sheet
            char_paths = char_candidate_paths[:needed]
            all_images.extend(char_paths)
            for j in range(len(char_paths)):
                on_progress(f"{prompt}: image {j + 1}/{len(char_paths)}")

            if options.double_sided:
                used_set = set(char_paths)
                available_back = [p for p in char_candidate_paths if p not in used_set]
                if len(available_back) >= needed:
                    char_back = available_back[:needed]
                elif char_paths:
                    char_back = available_back + [
                        char_paths[j % len(char_paths)]
                        for j in range(needed - len(available_back))
                    ]
                else:
                    char_back = available_back[:]
                all_back_images.extend(char_back)

            if not char_candidate_paths:
                dry_prompts.add(prompt)
            character_results.append(CharacterResult(
                prompt=prompt, image_paths=char_paths, candidate_paths=char_candidate_paths
            ))

        # Redistribute unfilled slots round-robin across all prompts
        shortfall = total - len(all_images)
        if shortfall > 0:
            for cr in character_results * shortfall:
                if len(all_images) >= total:
                    break
                if cancel_event and cancel_event.is_set():
                    raise InterruptedError("Cancelled by user.")
                # Skip prompts that already came up empty — re-searching them just
                # returns the same URLs (all blocked by `seen`) and wastes API calls.
                if cr.prompt in dry_prompts:
                    continue
                on_progress(f"Filling slot {len(all_images) + 1}/{total}: {cr.prompt}...")
                # Offset past the results already consumed so we fetch fresh URLs.
                extra_urls = search_images(
                    cr.prompt,
                    10,
                    start=random.randint(10, 30),
                    engine=options.search_engine
                )
                with ThreadPoolExecutor(max_workers=4) as rex:
                    rfuts = {rex.submit(process_image, url, (700, 1000), seen, seen_lock, temp_dir): url for url in extra_urls}
                    for rf in as_completed(rfuts):
                        if len(all_images) >= total:
                            for f in rfuts: f.cancel()
                            break
                        path = rf.result()
                        if not path:
                            continue
                        all_images.append(path)
                        cr.image_paths.append(path)
                        for f in rfuts: f.cancel()
                        break

    except InterruptedError:
        raise PipelineCancelled(all_images, character_results)

    # Guaranteed fallback: cycle existing images to fill any remaining blank slots
    if all_images and len(all_images) < total:
        pool = list(all_images)
        needed = total - len(all_images)
        all_images += [pool[i % len(pool)] for i in range(needed)]
        on_progress(f"Padded {needed} blank slot(s) by repeating existing images.")

    # Nothing usable at all — surface a clear reason instead of saving a blank sheet
    if not all_images:
        raise RuntimeError(
            "No usable images were found for any character. Check your API key / "
            "search engine selection, or try different names."
        )

    # Pad back images to match front count (shortfall slots get cycled back images)
    if options.double_sided and all_images:
        if not all_back_images:
            all_back_images = list(all_images)
        while len(all_back_images) < len(all_images):
            all_back_images.append(all_back_images[len(all_back_images) % len(all_back_images)])
        all_back_images = all_back_images[:len(all_images)]

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    if output_dir is None:
        output_dir = os.path.join(os.path.expanduser("~"), "Downloads")
    os.makedirs(output_dir, exist_ok=True)
    output_docx = os.path.join(output_dir, make_sheet_filename(prompts, timestamp))

    back_paths_for_doc = None
    if options.double_sided and all_back_images:
        on_progress("Building back sheet...")
        back_paths_for_doc = _mirror_rows(
            all_back_images[:options.rows * options.cols], options.rows, options.cols
        )

    on_progress("Building document...")
    create_doc(all_images, output_docx, options.rows, options.cols, options.paper_size,
               back_paths=back_paths_for_doc)

    output_pdf = None
    if options.export_pdf:
        soffice = find_soffice()
        if not soffice:
            on_progress("Warning: OpenOffice/LibreOffice not found. PDF skipped.")
        else:
            on_progress("Converting to PDF...")
            try:
                subprocess.run(
                    [soffice, "--headless", "--convert-to", "pdf", "--outdir", output_dir, output_docx],
                    check=True
                )
                output_pdf = output_docx.replace(".docx", ".pdf")
                on_progress(f"PDF created: {os.path.basename(output_pdf)}")
            except Exception as e:
                on_progress(f"PDF conversion failed: {e}")

    if options.send_email_to:
        to_addr = options.send_email_to.strip()
        if to_addr:
            on_progress(f"Sending email to {to_addr}...")
            try:
                send_email(to_addr, output_docx)
                on_progress("Email sent.")
            except Exception as e:
                on_progress(f"Email failed: {e}")

    return PipelineResult(
        output_docx=output_docx,
        output_pdf=output_pdf,
        characters=character_results,
        back_image_paths=back_paths_for_doc or [],
    )
